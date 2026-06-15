import os
import json
import numpy as np
import matplotlib
matplotlib.use('Agg') # Non-interactive backend
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from PIL import Image
import cv2
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls

def draw_system_architecture():
    """Generate System Architecture Diagram as PNG."""
    print("Menggambar diagram arsitektur sistem...")
    fig, ax = plt.subplots(figsize=(8.5, 5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    
    fig.patch.set_facecolor('#ffffff')
    
    # Draw Frontend Box
    frontend_box = patches.FancyBboxPatch((0.5, 1.0), 3.5, 4.0, boxstyle="round,pad=0.2", facecolor='#f8fafc', edgecolor='#94a3b8', linewidth=1.5)
    ax.add_patch(frontend_box)
    ax.text(2.25, 5.2, "Frontend (Next.js - Port 3000)", fontsize=10, fontweight='bold', ha='center', color='#0f172a')
    
    # Frontend elements
    ax.add_patch(patches.FancyBboxPatch((0.8, 3.8), 2.9, 0.7, boxstyle="round,pad=0.1", facecolor='#fee2e2', edgecolor='#f87171', linewidth=1))
    ax.text(2.25, 4.15, "UI Dashboard & Deteksi\n(Tailwind & TypeScript)", fontsize=8, ha='center', color='#991b1b')
    
    ax.add_patch(patches.FancyBboxPatch((0.8, 2.4), 2.9, 0.7, boxstyle="round,pad=0.1", facecolor='#fee2e2', edgecolor='#f87171', linewidth=1))
    ax.text(2.25, 2.75, "React State & Hooks\n(Inference UI handler)", fontsize=8, ha='center', color='#991b1b')

    ax.add_patch(patches.FancyBboxPatch((0.8, 1.2), 2.9, 0.5, boxstyle="round,pad=0.1", facecolor='#e0f2fe', edgecolor='#38bdf8', linewidth=1))
    ax.text(2.25, 1.45, "Client Fetch API (Fetch)", fontsize=8, ha='center', color='#0369a1')

    # Draw Backend Box
    backend_box = patches.FancyBboxPatch((6.0, 1.0), 3.5, 4.0, boxstyle="round,pad=0.2", facecolor='#f8fafc', edgecolor='#94a3b8', linewidth=1.5)
    ax.add_patch(backend_box)
    ax.text(7.75, 5.2, "Backend (FastAPI - Port 8000)", fontsize=10, fontweight='bold', ha='center', color='#0f172a')
    
    # Backend elements
    ax.add_patch(patches.FancyBboxPatch((6.3, 3.8), 2.9, 0.7, boxstyle="round,pad=0.1", facecolor='#e0f2fe', edgecolor='#38bdf8', linewidth=1))
    ax.text(7.75, 4.15, "API Endpoints\n(FastAPI Server)", fontsize=8, ha='center', color='#0369a1')
    
    ax.add_patch(patches.FancyBboxPatch((6.3, 2.8), 2.9, 0.5, boxstyle="round,pad=0.1", facecolor='#f0fdf4', edgecolor='#4ade80', linewidth=1))
    ax.text(7.75, 3.05, "OpenCV & PIL Image Processor", fontsize=8, ha='center', color='#166534')

    ax.add_patch(patches.FancyBboxPatch((6.3, 1.8), 2.9, 0.5, boxstyle="round,pad=0.1", facecolor='#f0fdf4', edgecolor='#4ade80', linewidth=1))
    ax.text(7.75, 2.05, "TensorFlow & Keras Engine", fontsize=8, ha='center', color='#166534')

    # Draw Arrows
    # Arrow from Frontend Fetch to Backend API
    ax.annotate("", xy=(6.0, 1.45), xytext=(4.0, 1.45),
                arrowprops=dict(arrowstyle="->", color='#0284c7', lw=1.5))
    ax.text(5.0, 1.6, "HTTP Request\n(File/JSON)", fontsize=8, ha='center', color='#0369a1')
    
    # Arrow from Backend API back to Frontend Fetch
    ax.annotate("", xy=(4.0, 2.75), xytext=(6.0, 2.75),
                arrowprops=dict(arrowstyle="->", color='#22c55e', lw=1.5))
    ax.text(5.0, 2.9, "HTTP Response\n(JSON Base64)", fontsize=8, ha='center', color='#166534')

    plt.tight_layout()
    plt.savefig('system_architecture.png', dpi=200, bbox_inches='tight')
    plt.close()

def draw_sequence_diagram():
    """Generate Sequence Diagram as PNG."""
    print("Menggambar diagram sequence alur deteksi...")
    fig, ax = plt.subplots(figsize=(8.5, 5.5))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 6)
    ax.axis('off')
    
    fig.patch.set_facecolor('#ffffff')
    
    lifelines = [1.5, 4.0, 6.5, 9.0]
    labels = ["User (Next.js)", "FastAPI", "OpenCV", "Keras CNN"]
    colors = ["#ef4444", "#0284c7", "#22c55e", "#8b5cf6"]
    
    for x, label, color in zip(lifelines, labels, colors):
        # Draw header block
        ax.add_patch(patches.FancyBboxPatch((x-0.9, 5.1), 1.8, 0.45, boxstyle="round,pad=0.05", facecolor=color, edgecolor='none'))
        ax.text(x, 5.32, label, fontsize=9, fontweight='bold', color='white', ha='center')
        # Draw dotted lifeline
        ax.plot([x, x], [0.5, 5.1], color='#cbd5e1', linestyle='--', linewidth=1)
        
    # Sequence arrows
    # 1. User -> API: POST /api/predict
    ax.annotate("", xy=(4.0, 4.5), xytext=(1.5, 4.5), arrowprops=dict(arrowstyle="->", color='#0f172a', lw=1.2))
    ax.text(2.75, 4.6, "1. POST /predict (Image File)", fontsize=7.5, ha='center', color='#0f172a', fontweight='semibold')
    
    # 2. API -> OpenCV: process (Grayscale, HSV, Canny)
    ax.annotate("", xy=(6.5, 3.9), xytext=(4.0, 3.9), arrowprops=dict(arrowstyle="->", color='#0284c7', lw=1.2))
    ax.text(5.25, 4.0, "2. Process CG Filters", fontsize=7.5, ha='center', color='#0369a1')
    
    # 3. OpenCV -> API: Base64 images
    ax.annotate("", xy=(4.0, 3.4), xytext=(6.5, 3.4), arrowprops=dict(arrowstyle="->", color='#22c55e', lw=1.2, linestyle=':'))
    ax.text(5.25, 3.5, "3. Return CG Images (Base64)", fontsize=7.5, ha='center', color='#166534')

    # 4. API -> CNN: Predict & Extract Maps
    ax.annotate("", xy=(9.0, 2.8), xytext=(4.0, 2.8), arrowprops=dict(arrowstyle="->", color='#0284c7', lw=1.2))
    ax.text(6.5, 2.9, "4. Predict & Activation Maps (batch)", fontsize=7.5, ha='center', color='#0369a1')

    # 5. CNN -> API: Scores & Activation Maps
    ax.annotate("", xy=(4.0, 2.2), xytext=(9.0, 2.2), arrowprops=dict(arrowstyle="->", color='#8b5cf6', lw=1.2, linestyle=':'))
    ax.text(6.5, 2.3, "5. Return CNN Predictions & Maps", fontsize=7.5, ha='center', color='#6d28d9')

    # 6. API -> User: Response JSON
    ax.annotate("", xy=(1.5, 1.4), xytext=(4.0, 1.4), arrowprops=dict(arrowstyle="->", color='#0f172a', lw=1.2, linestyle=':'))
    ax.text(2.75, 1.5, "6. Return JSON response (Base64)", fontsize=7.5, ha='center', color='#0f172a', fontweight='semibold')

    plt.tight_layout()
    plt.savefig('sequence_diagram.png', dpi=200, bbox_inches='tight')
    plt.close()

def main():
    print("Memulai pembuatan dokumen Word (.docx) lengkap dengan diagram dan hasil uji...")
    
    # 1. Generate Diagram Images
    try:
        draw_system_architecture()
        draw_sequence_diagram()
    except Exception as e:
        print(f"Error drawing diagrams: {str(e)}")
        
    doc = Document()
    
    # Document title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Laporan Alur Program: Sistem Deteksi Kematangan Tomat\n(CNN & Grafika Komputer)")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(180, 20, 20) # Tomato Red
    
    # Author / Meta info
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta_p.add_run("Laporan Tugas Akhir / Studi Kasus Grafika Komputer & Pengolahan Citra Digital\nTeknologi: Next.js (Frontend) & FastAPI (Backend) & Keras (CNN)")
    meta_run.font.name = 'Arial'
    meta_run.font.size = Pt(11)
    meta_run.font.italic = True
    meta_run.font.color.rgb = RGBColor(100, 100, 100)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    # Heading helpers
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59)
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(14)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = RGBColor(71, 85, 105)
        return h

    # Section 1
    add_heading_1("1. Arsitektur Sistem (Decoupled Client-Server)")
    
    p = doc.add_paragraph("Sistem dibangun menggunakan arsitektur modern yang memisahkan secara penuh antara antarmuka pengguna (Frontend) dengan mesin komputasi cerdas (Backend):")
    p.paragraph_format.space_after = Pt(8)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Frontend (Next.js - Port 3000): ").bold = True
    p.add_run("Menggunakan UI berbasis web modern (React, TypeScript, TailwindCSS) untuk menampilkan dashboard teoretis, generator dataset, log training model secara live, dan analisis deteksi citra.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Backend (FastAPI - Port 8000): ").bold = True
    p.add_run("Server Python untuk menangani inferensi model CNN (Keras/TensorFlow), pemrosesan citra digital dengan OpenCV (Grayscale, HSV Hue, Canny Edges), pembuatan dataset sintetis, dan proses pelatihan model di background task.")
    
    # Embed System Architecture Image
    if os.path.exists("system_architecture.png"):
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture("system_architecture.png", width=Inches(6.2))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(12)
        c_run = caption_p.add_run("Gambar 1.1 Diagram Arsitektur Decoupled Client-Server (Next.js & FastAPI)")
        c_run.font.size = Pt(9)
        c_run.font.italic = True
        c_run.font.color.rgb = RGBColor(100, 100, 100)

    # Section 2
    add_heading_1("2. Alur Program Utama (Main Program Flow)")
    
    doc.add_paragraph("Program memiliki tiga alur kerja utama yang dapat dijalankan secara interaktif oleh pengguna:")
    
    add_heading_2("A. Alur Pembuatan Dataset Sintetis (Procedural Rendering)")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Permintaan Pengguna: ").bold = True
    p.add_run("Pengguna memasukkan jumlah gambar per kategori (contoh: 100) di tab Dataset Generator Next.js dan menekan tombol Generate.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Kirim Parameter: ").bold = True
    p.add_run("Frontend Next.js mengirimkan request POST ke ")
    p.add_run("http://localhost:8000/api/generate-dataset").italic = True
    p.add_run(" berisi parameter JSON.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Pembersihan Direktori: ").bold = True
    p.add_run("Backend menghapus data lama dan menyiapkan direktori dataset/mentah/, dataset/matang/, dan dataset/busuk/.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Rendering Prosedural (Grafika Komputer): ").bold = True
    p.add_run("Setiap citra tomat digambar secara prosedural dengan:")
    
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("Latar Belakang: ").bold = True
    p.add_run("Render abu-abu gelap dengan tambahan noise grain.")
    
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("Tubuh Tomat: ").bold = True
    p.add_run("Menggambar lingkaran konsentris bertumpuk dengan gradasi warna sesuai kelasnya (Hijau untuk Mentah, Merah untuk Matang, Cokelat/Kuning kusam untuk Busuk) untuk efek volume 3D.")
    
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("Bercak Busuk (Khusus Kelas Busuk): ").bold = True
    p.add_run("Menggambar elips cokelat gelap dengan filter Gaussian Blur lembut untuk mensimulasikan kebusukan alami.")
    
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("Efek Kilau (Specular Highlight): ").bold = True
    p.add_run("Menggambar elips putih transparan miring 25 derajat (dibuat redup/pudar pada kelas busuk).")
    
    p = doc.add_paragraph(style='List Bullet 2')
    p.add_run("Tangkai Daun (Stem): ").bold = True
    p.add_run("Menggambar poligon berbentuk bintang hijau di bagian tengah atas tomat.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Penyimpanan: ").bold = True
    p.add_run("Backend menyimpan gambar dalam format .jpg ke subfolder masing-masing secara otomatis.")

    add_heading_2("B. Alur Pelatihan Model CNN (Model Training)")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Pemicu Pelatihan: ").bold = True
    p.add_run("Pengguna menekan tombol Mulai Pelatihan Model di tab Pelatihan Next.js.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Proses Asinkronus: ").bold = True
    p.add_run("Backend menerima request POST /api/train-model dan menjalankannya sebagai Background Task agar tidak menahan response web.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Partisi Data: ").bold = True
    p.add_run("Membaca citra dari dataset/ dan membaginya menjadi 80% Data Latih dan 20% Data Validasi.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Pelatihan Model: ").bold = True
    p.add_run("Keras melatih CNN sebanyak 10 epoch. Log pelatihan epoch dikirim ke state logs global di backend.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Polling Status: ").bold = True
    p.add_run("Next.js melakukan request GET /api/train-status setiap 1,5 detik untuk memperbarui terminal konsol log dan progress bar pelatihan.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Hasil Evaluasi: ").bold = True
    p.add_run("Menyimpan model ke model_tomat.h5 dan menggambar kurva akurasi/loss ke training_history.png.")

    add_heading_2("C. Alur Deteksi Kematangan & Analisis Citra (Inference & CG Pipeline)")
    
    p = doc.add_paragraph("Alur pendeteksian kematangan tomat yang diunggah oleh pengguna diproses secara runtut sesuai dengan diagram sekuensial berikut:")
    
    # Embed Sequence Diagram Image
    if os.path.exists("sequence_diagram.png"):
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture("sequence_diagram.png", width=Inches(6.2))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(12)
        c_run = caption_p.add_run("Gambar 2.1 Diagram Sekuensial Proses Deteksi & Inferensi Model")
        c_run.font.size = Pt(9)
        c_run.font.italic = True
        c_run.font.color.rgb = RGBColor(100, 100, 100)

    p = doc.add_paragraph(style='List Number')
    p.add_run("Input Gambar: ").bold = True
    p.add_run("Pengguna mengunggah gambar tomat kustom atau memilih tombol acak uji dataset.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Kirim Gambar: ").bold = True
    p.add_run("Frontend Next.js mengirimkan data form-data POST /api/predict ke FastAPI.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Mekanisme Grafika Komputer (OpenCV): ").bold = True
    p.add_run("Backend memproses gambar mentah secara paralel menjadi 3 representasi citra digital klasik: Grayscale, HSV Hue, dan Canny Edge (siluet).")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Inferensi Model CNN: ").bold = True
    p.add_run("Meresize gambar ke 128x128 piksel dan mengumpankannya ke model.predict() untuk menghitung nilai probabilitas kelas.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Ekstraksi Feature Maps: ").bold = True
    p.add_run("Mengambil aktivasi tensor dari lapisan konvolusi pertama (conv_1), mengambil 8 filter teratas, memetakan ke colormap Viridis, dan menyimpannya ke format base64 PNG.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Response JSON: ").bold = True
    p.add_run("Mengembalikan JSON berisi label terprediksi, keyakinan %, histogram probabilitas, 4 gambar tahapan citra, dan 8 gambar filter CNN.")

    # Section 3
    add_heading_1("3. Detail Mekanisme Pengolahan Citra Digital (Grafika Komputer)")
    
    p = doc.add_paragraph("Pada tahap inferensi citra, sistem memproses citra masukan menggunakan pustaka OpenCV:")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Grayscale (Luminositas): ").bold = True
    p.add_run("Mengubah piksel RGB menjadi intensitas cahaya tunggal Y menggunakan rumus standar ITU-R BT.601: ")
    p.add_run("Y = 0.299*R + 0.587*G + 0.114*B").italic = True
    p.add_run(". Langkah ini berguna untuk menghilangkan derau warna dan berfokus pada kecerahan.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("HSV Hue Channel (Warna Murni): ").bold = True
    p.add_run("Mengonversi RGB ke HSV (Hue, Saturation, Value). Dengan memisahkan saluran Hue (warna murni), sistem dapat mengidentifikasi area hijau (tomat mentah), merah (tomat matang), dan cokelat kusam (tomat busuk) tanpa terdistorsi oleh intensitas bayangan cahaya.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Canny Edge Detection (Deteksi Kontur): ").bold = True
    p.add_run("Mendeteksi garis tepi siluet tomat. Menggunakan filter Gaussian untuk mengurangi noise, operator gradien Sobel spasial untuk menghitung transisi kecerahan, Non-Maximum Suppression untuk menipiskan garis tepi, serta Hysteresis Thresholding (ambang batas bawah 50 dan atas 150) untuk menarik kontur tomat yang tegas.")

    # Section 4
    add_heading_1("4. Detail Struktur & Alur CNN (Convolutional Neural Network)")
    
    p = doc.add_paragraph("Setelah citra diubah dimensinya menjadi 128x128 piksel, citra diumpankan ke model CNN dengan alur sebagai berikut:")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Lapisan Rescaling: ").bold = True
    p.add_run("Mengubah nilai piksel [0, 255] menjadi [0.0, 1.0] dengan mengalikannya dengan 1/255.0.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Blok Konvolusi Pertama (conv_1): ").bold = True
    p.add_run("Menerapkan 16 filter 3x3 piksel untuk mengekstrak fitur tingkat dasar seperti batas garis tepi, sudut, dan gradasi warna primer.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Max Pooling: ").bold = True
    p.add_run("Menggunakan jendela 2x2 piksel untuk mengambil nilai maksimal, mereduksi dimensi spasial menjadi setengahnya (64x64) guna menghemat komputasi.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Blok Konvolusi 2 & 3: ").bold = True
    p.add_run("Menggunakan masing-masing 32 dan 64 filter untuk mengekstrak fitur yang lebih tinggi dan abstrak seperti pola tektur kulit dan bentuk kebulatan tomat.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Flattening & Dropout: ").bold = True
    p.add_run("Meratakan matriks fitur menjadi vektor 1D. Menerapkan Dropout (0.2) untuk menonaktifkan secara acak 20% neuron guna mencegah overfitting selama pelatihan.")
    
    p = doc.add_paragraph(style='List Number')
    p.add_run("Fully Connected & Softmax: ").bold = True
    p.add_run("Menggunakan lapisan Dense dengan fungsi aktivasi ReLU, diakhiri dengan lapisan output Softmax (3 unit) untuk memproyeksikan nilai probabilitas klasifikasi kematangan tomat (Mentah, Busuk, Matang) yang bernilai total 100%.")

    # ==================== NEW SECTION: TEST CASES ====================
    add_heading_1("5. Hasil Pengujian Model (Uji Kasus Gambar Tomat Fisik)")
    
    p = doc.add_paragraph("Bagian ini menampilkan hasil pengujian sistem terhadap 3 buah gambar tomat fisik (TOMAT1.jpeg, TOMAT2.jpeg, TOMAT3.jpg) yang disediakan untuk menguji performa model CNN dalam mengklasifikasikan kondisi kematangan tomat secara real-time.")
    
    import tensorflow as tf
    from tensorflow.keras.models import Model
    
    # Load model and class mapping
    try:
        model = tf.keras.models.load_model("model_tomat.h5")
        with open("class_mapping.json", "r") as f:
            mapping = json.load(f)
            class_labels = [mapping[str(idx)] for idx in range(len(mapping))]
            
        conv1_layer = model.get_layer('conv_1')
        activation_model = Model(inputs=model.inputs, outputs=conv1_layer.output)
    except Exception as e:
        print(f"Error loading model for test cases: {str(e)}")
        model = None
        
    test_files = ["TOMAT1.jpeg", "TOMAT2.jpeg", "TOMAT3.jpg"]
    
    for i, filename in enumerate(test_files):
        if not os.path.exists(filename):
            print(f"File {filename} tidak ditemukan, skip.")
            continue
            
        print(f"Menganalisis file uji {filename}...")
        
        # Open image
        image = Image.open(filename).convert("RGB")
        img_np = np.array(image)
        
        # Preprocessing
        gray = cv2.cvtColor(img_np, cv2.COLOR_RGB2GRAY)
        hsv = cv2.cvtColor(img_np, cv2.COLOR_RGB2HSV)
        h_channel = hsv[:, :, 0]
        edges = cv2.Canny(gray, 50, 150)
        
        img_resized = image.resize((128, 128))
        img_resized_np = np.array(img_resized)
        
        # Model Prediction
        pred_label = "Error"
        confidence_pct = 0.0
        preds_dict = {}
        
        if model is not None:
            img_batch = np.expand_dims(img_resized_np, axis=0)
            preds = model.predict(img_batch)[0]
            pred_idx = np.argmax(preds)
            pred_label = class_labels[pred_idx]
            confidence_pct = float(preds[pred_idx]) * 100
            preds_dict = {class_labels[idx]: float(preds[idx]) * 100 for idx in range(len(class_labels))}
            
            # Plot intermediate image processing stages
            fig, axes = plt.subplots(1, 4, figsize=(10, 2.5))
            axes[0].imshow(img_np)
            axes[0].set_title("Original (Warna)", fontsize=8)
            axes[0].axis("off")
            
            axes[1].imshow(gray, cmap="gray")
            axes[1].set_title("Grayscale (Luminositas)", fontsize=8)
            axes[1].axis("off")
            
            axes[2].imshow(h_channel, cmap="hsv")
            axes[2].set_title("Hue (Warna Murni)", fontsize=8)
            axes[2].axis("off")
            
            axes[3].imshow(edges, cmap="gray")
            axes[3].set_title("Canny Edge (Kontur)", fontsize=8)
            axes[3].axis("off")
            
            plt.tight_layout()
            pipeline_filename = f"test_case_{i}_pipeline.png"
            plt.savefig(pipeline_filename, dpi=150, bbox_inches='tight')
            plt.close()
            
            # Plot CNN Activation Maps
            activations = activation_model.predict(img_batch)[0]
            fig, axes = plt.subplots(2, 4, figsize=(8, 4))
            for f_idx in range(8):
                ax = axes[f_idx // 4, f_idx % 4]
                f_map = activations[:, :, f_idx]
                f_min, f_max = f_map.min(), f_map.max()
                if f_max > f_min:
                    f_map = (f_map - f_min) / (f_max - f_min)
                ax.imshow(f_map, cmap="viridis")
                ax.set_title(f"Filter {f_idx+1}", fontsize=7)
                ax.axis("off")
                
            plt.tight_layout()
            activation_filename = f"test_case_{i}_activations.png"
            plt.savefig(activation_filename, dpi=150, bbox_inches='tight')
            plt.close()
            
        # Add test case sub-heading in docx
        add_heading_2(f"Uji Kasus {i+1}: Berkas {filename}")
        
        # Test Case Details
        p_desc = doc.add_paragraph()
        p_desc.add_run(f"Nama File: ").bold = True
        p_desc.add_run(f"{filename}\n")
        p_desc.add_run(f"Hasil Klasifikasi Model CNN: ").bold = True
        
        # Color coding class name text
        run_class = p_desc.add_run(f"{pred_label.upper()}")
        run_class.bold = True
        if pred_label == "matang":
            run_class.font.color.rgb = RGBColor(180, 20, 20)
        elif pred_label == "mentah":
            run_class.font.color.rgb = RGBColor(16, 120, 16)
        else:
            run_class.font.color.rgb = RGBColor(180, 100, 20)
            
        p_desc.add_run(f" (Tingkat Keyakinan: ").italic = True
        p_desc.add_run(f"{confidence_pct:.2f}%").bold = True
        p_desc.add_run(f")\n").italic = True
        
        p_desc.add_run("Distribusi Probabilitas Kelas:\n").bold = True
        for label_key, val_pct in preds_dict.items():
            p_desc.add_run(f"  • {label_key.capitalize()}: ").italic = True
            p_desc.add_run(f"{val_pct:.2f}%\n").bold = True
            
        # Embed Original & Pipeline Image
        pipeline_filename = f"test_case_{i}_pipeline.png"
        if os.path.exists(pipeline_filename):
            doc.add_paragraph("Visualisasi Tahapan Pengolahan Citra Klasik:").paragraph_format.space_after = Pt(2)
            img_p = doc.add_paragraph()
            img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_p.add_run().add_picture(pipeline_filename, width=Inches(5.8))
            img_p.paragraph_format.space_after = Pt(6)
            
            # Embed Activations Image
            activation_filename = f"test_case_{i}_activations.png"
            if os.path.exists(activation_filename):
                doc.add_paragraph("Peta Aktivasi Konvolusi CNN Pertama (Feature Maps):").paragraph_format.space_after = Pt(2)
                img_p = doc.add_paragraph()
                img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_p.add_run().add_picture(activation_filename, width=Inches(5.0))
                img_p.paragraph_format.space_after = Pt(14)

    # Section 6: Antarmuka Pengguna (Web User Interface)
    add_heading_1("6. Antarmuka Pengguna (Web User Interface)")
    
    p = doc.add_paragraph(
        "Aplikasi ini dilengkapi dengan antarmuka pengguna (UI) berbasis web yang responsif dan modern menggunakan "
        "Next.js. Desain web dirancang menggunakan tema gelap premium (Dark Mode) untuk memberikan tampilan yang "
        "estetis, mewah, dan profesional. Berikut adalah screenshot dari antarmuka web aplikasi deteksi kematangan tomat:"
    )
    p.paragraph_format.space_after = Pt(8)
    
    if os.path.exists("web_ui_mockup.png"):
        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture("web_ui_mockup.png", width=Inches(6.0))
        caption_p = doc.add_paragraph()
        caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_p.paragraph_format.space_after = Pt(12)
        c_run = caption_p.add_run("Gambar 6.1 Tampilan Antarmuka Web Dashboard Sistem Deteksi Kematangan Tomat")
        c_run.font.size = Pt(9)
        c_run.font.italic = True
        c_run.font.color.rgb = RGBColor(100, 100, 100)
        
    p = doc.add_paragraph("Fitur-fitur utama yang tersedia pada antarmuka web ini meliputi:")
    p.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Dashboard Analitis: ").bold = True
    p.add_run("Menampilkan informasi ringkasan status model, akurasi pelatihan, serta navigasi cepat ke fitur utama.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Dataset Generator: ").bold = True
    p.add_run("Form interaktif untuk mengatur jumlah citra tomat sintetis yang ingin dihasilkan secara prosedural untuk kelas Mentah, Matang, dan Busuk.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Live Training Console: ").bold = True
    p.add_run("Terminal log interaktif yang menampilkan proses pelatihan model CNN epoch demi epoch secara real-time, lengkap dengan grafik kurva akurasi dan loss yang terupdate otomatis.")
    
    p = doc.add_paragraph(style='List Bullet')
    p.add_run("Real-time Inference: ").bold = True
    p.add_run("Area drag-and-drop untuk mengunggah berkas gambar tomat kustom atau melakukan uji acak. Hasil pengujian akan menampilkan probabilitas kelas dalam bentuk progress bar berwarna, visualisasi pemrosesan citra digital klasik (Grayscale, HSV, Canny), dan grid visualisasi filter konvolusi CNN (Feature Maps).")

    # Save Document
    filename = "laporan_alur_program_final.docx"
    try:
        doc.save(filename)
        print(f"Berhasil membuat dokumen Word: {filename}")
    except PermissionError:
        print(f"Warning: {filename} sedang dibuka oleh aplikasi lain (locked). Mencoba nama file alternatif...")
        filename = "laporan_alur_program_final_v2.docx"
        try:
            doc.save(filename)
            print(f"Berhasil membuat dokumen Word (alternatif): {filename}")
        except Exception as ex:
            import time
            filename = f"laporan_alur_program_final_{int(time.time())}.docx"
            doc.save(filename)
            print(f"Berhasil membuat dokumen Word (timestamp): {filename}")
    
    # Cleanup temporary images
    for i in range(len(test_files)):
        pipeline_filename = f"test_case_{i}_pipeline.png"
        activation_filename = f"test_case_{i}_activations.png"
        if os.path.exists(pipeline_filename):
            os.remove(pipeline_filename)
        if os.path.exists(activation_filename):
            os.remove(activation_filename)

if __name__ == "__main__":
    main()

import os
import json
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, color_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_paragraph_with_spacing(doc, text="", style=None, space_before=0, space_after=6, line_spacing=1.15):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    return p

def main():
    print("Memulai pembuatan dokumen Word (.docx) Catatan Ujian...")
    doc = Document()
    
    # Page setup - Margins (1 inch or 2.54 cm)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(4)
    title_run = title_p.add_run("CATATAN STUDI & CHEAT SHEET UJIAN")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(180, 20, 20) # Tomato Red
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    sub_run = subtitle_p.add_run("Sistem Deteksi Kematangan Tomat\n(Grafika Komputer Prosedural, Pengolahan Citra Digital, & CNN)")
    sub_run.font.name = 'Arial'
    sub_run.font.size = Pt(12)
    sub_run.font.italic = True
    sub_run.font.color.rgb = RGBColor(100, 100, 100)

    # Heading helpers
    def add_heading_1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(18)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 41, 59) # Slate 800
        return h

    def add_heading_2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(71, 85, 105) # Slate 600
        return h

    def add_body(text, bold_prefix=None, italic=False, space_after=6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            br = p.add_run(bold_prefix)
            br.font.name = 'Arial'
            br.font.size = Pt(10.5)
            br.bold = True
            br.font.color.rgb = RGBColor(51, 65, 85)
            
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        r.italic = italic
        return p

    def add_bullet(text, bold_prefix=None, space_after=4):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        
        if bold_prefix:
            br = p.add_run(bold_prefix)
            br.font.name = 'Arial'
            br.font.size = Pt(10.5)
            br.bold = True
            br.font.color.rgb = RGBColor(51, 65, 85)
            
        r = p.add_run(text)
        r.font.name = 'Arial'
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(51, 65, 85)
        return p

    # --- SECTION I ---
    add_heading_1("I. Arsitektur & Identitas Sistem (Decoupled Client-Server)")
    add_body("Sistem dibangun dengan memisahkan secara penuh antarmuka pengguna dengan mesin komputasi untuk efisiensi beban kerja:")
    add_bullet("React, TypeScript, dan TailwindCSS. Menyediakan antarmuka dashboard teoretis, generator dataset, log training model secara live, dan analisis deteksi citra.", bold_prefix="Frontend (Next.js - Port 3000): ")
    add_bullet("Python, OpenCV, TensorFlow/Keras, dan PIL (Pillow). Server API untuk menangani inferensi model CNN, pemrosesan citra digital, pembuatan dataset prosedural acak, dan proses pelatihan model di background task.", bold_prefix="Backend (FastAPI - Port 8000): ")
    add_bullet("Klien mengirim data form-data (berkas citra) atau parameter JSON ke backend, lalu menerima respons JSON berisi probabilitas kelas dan citra tahapan pemrosesan dalam format string Base64 PNG.", bold_prefix="Mekanisme Komunikasi: ")

    # --- SECTION II ---
    add_heading_1("II. Aspek Grafika Komputer: Pembuatan Dataset Prosedural (Procedural Rendering)")
    add_body("Untuk mengatasi keterbatasan data fisik, sistem merender citra tomat secara dinamis menggunakan formula matematika:")
    
    add_heading_2("1. Representasi Geometri & Teknik Shading 3D")
    add_bullet("Tubuh tomat digambar menggunakan 15 tingkat lingkaran konsentris bertumpuk dari ukuran radius terbesar ke terkecil (1.0 hingga 0.4 x r).", bold_prefix="Volume 3D Sederhana (Spherical Shading): ")
    add_bullet("Pusat setiap lingkaran digeser secara bertahap sebesar 18% ke arah kiri-atas (cx - offset, cy - offset). Pergeseran ini mensimulasikan datangnya arah cahaya (light source direction) dari sudut kiri-atas.", bold_prefix="Arah Sumber Cahaya (Light Source Offset): ")
    add_bullet("Warna diinterpolasi secara linier dari tepi luar (lebih gelap/bayangan) ke arah tengah (lebih terang/cahaya menyebar) sesuai kategori kematangan:", bold_prefix="Interpolasi Gradasi Warna: ")
    add_bullet("Hijau gelap ke kuning-hijau terang.", bold_prefix="  • Mentah: ")
    add_bullet("Merah gelap ke merah-jingga terang.", bold_prefix="  • Matang: ")
    add_bullet("Menggunakan warna dasar yang dikalikan faktor redup (0.65 - 0.85) untuk memberikan tampilan kusam.", bold_prefix="  • Busuk: ")

    add_heading_2("2. Fitur Spesifik Objek & Transparansi (Alpha Blending)")
    add_bullet("Mensimulasikan pelapukan dengan menggambar 2 sampai 4 elips cokelat gelap acak, diblending dengan filter Gaussian Blur (radius 1.8) agar menyatu alami dengan kulit tomat.", bold_prefix="Bercak Busuk (Rot Spots): ")
    add_bullet("Menggambar elips putih transparan miring 25 derajat di bagian kiri atas tubuh tomat untuk efek pantulan cahaya.", bold_prefix="Pantulan Kilap (Specular Highlight): ")
    add_bullet("Nilai alpha diatur ke 140 (cukup mengkilap/segar).", bold_prefix="  - Kelas Matang & Mentah: ")
    add_bullet("Nilai alpha diatur ke 35 (sangat redup/kusam karena kerutan pembusukan).", bold_prefix="  - Kelas Busuk: ")
    add_bullet("Poligon berbentuk bintang hijau dengan 4 hingga 6 kelopak yang digambar di bagian tengah atas (muncul pada 70% gambar). Posisi kelopak dihitung dengan koordinat polar: x = cx + l * cos(theta), y = cy + l * sin(theta).", bold_prefix="Tangkai Daun (Stem): ")

    add_heading_2("3. Efek Tekstur & Post-processing")
    add_bullet("Menambahkan variasi nilai pixel acak ([-12, 12] untuk background, [-4, 4] untuk final) guna menyimulasikan sensor noise kamera asli.", bold_prefix="Noise Grain: ")
    add_bullet("Overlay hitam transparan di sudut-sudut gambar berbasis jarak Euclidean pusat untuk mensimulasikan degradasi pencahayaan tepi lensa kamera.", bold_prefix="Vignetting: ")

    # --- SECTION III ---
    add_heading_1("III. Aspek Pengolahan Citra Digital (PCD) Klasik pada Deteksi")
    add_body("Citra masukan diproses secara bertahap menggunakan OpenCV sebelum masuk ke neural network:")
    
    add_heading_2("1. Grayscale Conversion (Luminositas)")
    add_body("Mengonversi citra RGB (3 channel) menjadi Grayscale (1 channel) untuk menghilangkan informasi warna yang berlebihan, memfokuskan pada transisi kecerahan, dan meminimalkan noise warna.")
    add_body("Y = 0.299 * R + 0.587 * G + 0.114 * B", bold_prefix="Rumus ITU-R BT.601: ", italic=True)

    add_heading_2("2. HSV Color Space Conversion (Hue Channel)")
    add_body("Model RGB sangat peka terhadap perubahan intensitas pencahayaan (bayangan atau sorotan lampu). Dengan mengonversi ke HSV (Hue, Saturation, Value) dan mengambil saluran Hue (Warna Murni), sistem dapat mengidentifikasi piksel hijau (mentah), merah (matang), dan cokelat (busuk) secara konsisten tanpa terdistorsi oleh fluktuasi bayangan.")

    add_heading_2("3. Canny Edge Detection (Deteksi Kontur Tepi)")
    add_body("Mengekstraksi garis keliling objek melalui empat tahapan utama:")
    add_bullet("Menghilangkan derau piksel halus menggunakan Gaussian Filter.", bold_prefix="1. Noise Reduction: ")
    add_bullet("Menggunakan filter Sobel horizontal (Gx) dan vertikal (Gy) untuk mencari kekuatan (magnitudo) dan arah perubahan gradien kecerahan.", bold_prefix="2. Gradient Calculation: ")
    add_bullet("Menyeleksi piksel yang memiliki nilai lokal maksimum sepanjang arah gradien dan menekan piksel lainnya menjadi 0 untuk menipiskan garis tepi menjadi 1 piksel.", bold_prefix="3. Non-Maximum Suppression (NMS): ")
    add_bullet("Menggunakan ambang batas ganda (Lower = 50, Upper = 150). Piksel gradien > 150 adalah tepi kuat. Piksel < 50 dibuang. Piksel 50-150 adalah tepi lemah dan hanya dipertahankan jika terhubung dengan tepi kuat.", bold_prefix="4. Hysteresis Thresholding: ")

    # --- SECTION IV ---
    add_heading_1("IV. Arsitektur Jaringan Saraf Konvolusional (CNN)")
    add_body("Model CNN dirancang secara Sequential menggunakan pustaka TensorFlow/Keras:")
    
    add_heading_2("1. Normalisasi & Augmentasi Data")
    add_bullet("Membagi nilai piksel dengan 255.0 untuk menormalkan rentang [0, 255] menjadi [0.0, 1.0]. Ini mencegah gradien meledak (exploding gradient) dan mempercepat konvergensi pelatihan.", bold_prefix="Rescaling: ")
    add_bullet("Mencegah overfitting dengan menerapkan RandomFlip, RandomRotation(0.15), RandomTranslation(0.1, 0.1), dan RandomZoom(0.15) secara dinamis hanya saat pelatihan.", bold_prefix="Augmentation (Latih): ")

    add_heading_2("2. Struktur Lapisan Layer CNN")
    add_bullet("Menerapkan 16 filter berukuran 3x3 piksel dengan fungsi aktivasi ReLU untuk mengekstrak fitur tingkat dasar seperti garis tepi, transisi warna awal, dan sudut.", bold_prefix="Conv2D (conv_1): ")
    add_bullet("Mengambil nilai maksimal pada jendela 2x2, mereduksi dimensi spasial sebanyak 50% (dari 128x128 ke 64x64) guna menghemat komputasi dan memberikan kebalan pergeseran lokasi (translation invariance).", bold_prefix="MaxPooling2D: ")
    add_bullet("Menerapkan masing-masing 32 dan 64 filter untuk mempelajari pola visual abstrak seperti tekstur bercak busuk dan bentuk kelengkungan.", bold_prefix="Conv2D (conv_2 & conv_3): ")
    add_bullet("Meratakan peta fitur multi-dimensi menjadi vektor 1D agar dapat diproses oleh Dense Layer.", bold_prefix="Flatten: ")
    add_bullet("Lapisan Dense (64 unit, ReLU) dikombinasikan dengan Dropout (0.2) untuk menonaktifkan 20% neuron acak selama training demi mencegah ko-adaptasi neuron (overfitting).", bold_prefix="Fully Connected & Dropout: ")
    add_bullet("Lapisan output (3 unit) dengan fungsi Softmax untuk memproyeksikan skor numerik mentah menjadi nilai probabilitas kelas (Mentah, Busuk, Matang) dengan total akumulasi 100%.", bold_prefix="Output Layer (Softmax): ")

    add_heading_2("3. Hyperparameter Pelatihan")
    add_bullet("Adam (Adaptive Moment Estimation) - Laju pembelajaran adaptif.", bold_prefix="Optimizer: ")
    add_bullet("sparse_categorical_crossentropy - Karena label target berupa integer indeks kelas (0, 1, 2).", bold_prefix="Loss Function: ")
    add_bullet("80% Data Latih (Training) & 20% Data Validasi (Validation).", bold_prefix="Partisi Dataset: ")

    # --- SECTION V ---
    add_heading_1("V. Visualisasi Peta Aktivasi CNN (Feature Maps)")
    add_bullet("Representasi visual dari deteksi filter konvolusi spasial saat memproses citra.", bold_prefix="Definisi: ")
    add_bullet("Sistem mengekstrak output tensor dari lapisan konvolusi pertama ('conv_1'), mengambil 8 filter pertama, menormalisasi nilainya ke rentang [0, 255], lalu menerapkan colormap 'Viridis' (biru-kuning-hijau).", bold_prefix="Mekanisme: ")
    add_bullet("Filter tepi akan menampilkan garis terang di sekeliling tomat; filter warna mengaktifkan area daun hijau atau kulit merah; filter kecerahan bereaksi kuat pada specular highlight kiri atas.", bold_prefix="Interpretasi: ")

    # --- SECTION VI ---
    add_heading_1("VI. Ringkasan Tanya Jawab Cepat (Cheat Sheet Ujian)")
    
    # Add Table
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    
    # Table styling
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Pertanyaan Ujian'
    hdr_cells[1].text = 'Jawaban Cepat / Kunci Jawaban'
    
    # Width settings
    table.columns[0].width = Inches(2.5)
    table.columns[1].width = Inches(4.0)
    
    for cell in hdr_cells:
        set_cell_background(cell, "B41414") # Tomato red header
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        # Font style for header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Arial'
                r.font.size = Pt(10)
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)

    qa_list = [
        ("Bagaimana cara menyimulasikan efek 3D secara prosedural pada tomat?", 
         "Menggambar 15 lingkaran konsentris bertumpuk dari besar ke kecil dengan pergeseran titik pusat (offset) 18% ke kiri-atas untuk memodelkan datangnya arah cahaya (spherical shading)."),
        ("Mengapa ruang warna HSV lebih dipilih daripada RGB untuk segmentasi/klasifikasi warna?", 
         "Karena RGB sensitif terhadap perubahan intensitas pencahayaan. HSV memisahkan warna murni (Hue) dari kecerahan (Value), sehingga identifikasi warna tomat hijau/merah/cokelat tetap stabil meskipun berada di bawah bayangan atau kecerahan yang bervariasi."),
        ("Apa kegunaan Hysteresis Thresholding pada Canny Edge Detection?", 
         "Mengeliminasi tepi palsu menggunakan ambang batas ganda (50 dan 150). Piksel > 150 dianggap tepi kuat. Piksel < 50 dibuang. Piksel 50-150 dipertahankan hanya jika terhubung langsung dengan piksel tepi kuat."),
        ("Sebutkan rumus Grayscale yang digunakan beserta koefisiennya!", 
         "Menggunakan standar ITU-R BT.601: Y = 0.299 * R + 0.587 * G + 0.114 * B."),
        ("Apa fungsi dari Dropout layer dalam arsitektur CNN?", 
         "Mencegah overfitting dengan menonaktifkan secara acak 20% neuron pada Dense layer selama pelatihan agar neuron belajar mandiri dan tidak saling bergantung secara berlebihan."),
        ("Mengapa output layer CNN menggunakan Softmax?", 
         "Untuk memetakan skor mentah (logits) menjadi nilai probabilitas kelas (berkisar 0 s.d 1) dengan total akumulasi probabilitas seluruh kelas sama dengan 1.0 (100%)."),
        ("Bagaimana peta aktivasi (feature maps) divisualisasikan?", 
         "Mengambil output lapisan 'conv_1', menormalisasi nilai aktivasi ke rentang [0, 255], lalu mengubah matriks 2D tersebut ke citra berwarna menggunakan colormap Viridis."),
        ("Mengapa nilai piksel dibagi 255.0 (rescaling) sebelum masuk ke konvolusi?", 
         "Untuk menormalisasi rentang piksel [0, 255] menjadi floating point [0.0, 1.0]. Ini mencegah gradien meledak (exploding gradient) saat backpropagation dan mempercepat konvergensi optimasi.")
    ]

    for idx, (q, a) in enumerate(qa_list):
        row_cells = table.add_row().cells
        row_cells[0].text = q
        row_cells[1].text = a
        
        # Zebra striping background
        bg_color = "F8FAFC" if idx % 2 == 0 else "FFFFFF"
        
        for i, cell in enumerate(row_cells):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            
            # Set border or shading
            tcPr = cell._tc.get_or_add_tcPr()
            # Subtle border
            borders_xml = f'''
            <w:tcBorders {nsdecls("w")}>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                <w:top w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tcBorders>
            '''
            tcPr.append(parse_xml(borders_xml))
            
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                for r in p.runs:
                    r.font.name = 'Arial'
                    r.font.size = Pt(9.5)
                    r.font.color.rgb = RGBColor(51, 65, 85)
                    if i == 0:
                        r.font.bold = True
                        r.font.color.rgb = RGBColor(30, 41, 59)

    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    filename = "Catatan_Ujian_Grafika_Komputer.docx"
    doc.save(filename)
    print(f"Berhasil membuat dokumen Word: {filename}")

if __name__ == "__main__":
    main()
